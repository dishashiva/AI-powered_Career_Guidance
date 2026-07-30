"""Generate CareerAI project report as DOCX."""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION

BASE = os.path.dirname(os.path.abspath(__file__))
OUTPUT_DIR = os.path.join(BASE, "Outputs")
OUTPUT_DOCX = os.path.join(BASE, "CareerAI_Report.docx")

# ─── Colors ───────────────────────────────────────────────────
BLUE = RGBColor(0x2E, 0x6F, 0xF2)
DARK = RGBColor(0x10, 0x18, 0x28)
GRAY = RGBColor(0x47, 0x54, 0x67)
LIGHT_GRAY = RGBColor(0x98, 0xA2, 0xB3)

# ─── Setup ────────────────────────────────────────────────────
doc = Document()

# Page setup — A4
section = doc.sections[0]
section.page_width = Cm(21.0)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.0)
section.left_margin = Cm(2.54)
section.right_margin = Cm(2.54)
section.orientation = WD_ORIENTATION.PORTRAIT

# ─── Tune styles ──────────────────────────────────────────────
style_normal = doc.styles["Normal"]
style_normal.font.name = "Calibri"
style_normal.font.size = Pt(11)
style_normal.font.color.rgb = GRAY
style_normal.paragraph_format.line_spacing = 1.15
style_normal.paragraph_format.space_after = Pt(6)

style_title = doc.styles["Title"]
style_title.font.name = "Calibri Light"
style_title.font.size = Pt(28)
style_title.font.bold = True
style_title.font.color.rgb = DARK

style_subtitle = doc.styles["Subtitle"]
style_subtitle.font.name = "Calibri"
style_subtitle.font.size = Pt(14)
style_subtitle.font.color.rgb = GRAY

for level, size in [(1, 18), (2, 14), (3, 12)]:
    s = doc.styles[f"Heading {level}"]
    s.font.name = "Calibri Light"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = DARK
    s.paragraph_format.space_before = Pt(14 - 2 * level)
    s.paragraph_format.space_after = Pt(4)

style_caption = doc.styles["Caption"]
style_caption.font.name = "Calibri"
style_caption.font.size = Pt(9)
style_caption.font.italic = True
style_caption.font.color.rgb = LIGHT_GRAY


# ─── Helpers ──────────────────────────────────────────────────
def add_body(text):
    p = doc.add_paragraph(text, style="Normal")
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p

def add_image_with_caption(filename, caption):
    path = os.path.join(OUTPUT_DIR, filename)
    if not os.path.exists(path):
        add_body(f"[Screenshot missing: {filename}]")
        return
    # Scale to ~15cm width (fits within A4 margins)
    doc.add_picture(path, width=Cm(15))
    last_p = doc.paragraphs[-1]
    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption, style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph("CareerAI", style="Title")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("AI-Powered Career Intelligence Platform", style="Subtitle")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

p = doc.add_paragraph("INFOSYS SPRINGBOARD VIRTUAL INTERNSHIP")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.color.rgb = GRAY

p = doc.add_paragraph("Project: CareerAI — AI-Powered Career Intelligence Platform")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].font.size = Pt(12)
p.runs[0].font.color.rgb = GRAY

for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph("SUBMITTED BY:")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.runs[0].bold = True
p.runs[0].font.size = Pt(12)
p.runs[0].font.color.rgb = GRAY

for line in ["Name: Disha S", "Domain: Artificial Intelligence", "Date: 08.07.2026"]:
    p = doc.add_paragraph(line)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.color.rgb = GRAY

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)

add_body(
    'The project "CareerAI — AI-Powered Career Intelligence Platform" '
    'is a full-stack web application that functions as an intelligent career coach. '
    'It leverages artificial intelligence to analyze resumes, detect skill gaps, '
    'predict career paths, and provide real-time AI-powered guidance to help '
    'professionals make informed career decisions.'
)
add_body(
    'The platform addresses a common challenge faced by job seekers: understanding '
    'how their resume performs against applicant tracking systems (ATS), identifying '
    'which skills are missing for their target roles, and finding the right career '
    'trajectory. CareerAI automates this analysis using NLP-powered resume parsing '
    'and large language model inference, delivering actionable insights in seconds.'
)
add_body(
    'Users can upload their resume in PDF or DOCX format, receive an ATS compatibility '
    'score, view detected skills and skill gaps, explore AI-recommended career paths '
    'with salary estimates, browse personalized job and course recommendations, and '
    'interact with an AI career coach chatbot that understands their profile context.'
)

# ══════════════════════════════════════════════════════════════
# 2. OBJECTIVE
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. Objective", level=1)
add_body("The primary objectives of the CareerAI platform are:")

objectives = [
    ("<b>Resume Analysis & ATS Scoring</b> — Parse uploaded resumes using NLP techniques, extract structured data (skills, roles, experience), and generate an ATS compatibility score benchmarked against industry standards."),
    ("<b>Skill Gap Detection</b> — Identify missing skills relative to the user's target career path and provide prioritized recommendations with explanations for each gap."),
    ("<b>Career Path Prediction</b> — Use AI to generate personalized career trajectory recommendations with match percentages, required skills, salary ranges, and timelines."),
    ("<b>Salary Intelligence</b> — Provide role-aware, location-aware salary estimates based on the user's detected skillset and experience level."),
    ("<b>Job & Course Recommendations</b> — Curate personalized job listings and skill-building courses matched to the user's profile and identified gaps."),
    ("<b>AI Career Coach</b> — Offer a conversational chatbot interface that maintains profile context and delivers personalized, actionable career guidance."),
    ("<b>Secure Authentication</b> — Implement JWT-based user authentication with registration, login, and profile management."),
    ("<b>Clean, Professional UI</b> — Deliver a responsive, accessible interface built with React that prioritizes usability and visual clarity."),
]

for obj in objectives:
    add_bullet(obj)

# ══════════════════════════════════════════════════════════════
# 3. TECH STACK
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. Tech Stack", level=1)
add_body("The CareerAI platform was built using the following technologies:")

tech_data = [
    ("Frontend", "React.js + Vite", "Single-page application with fast dev server and optimized builds"),
    ("Backend", "FastAPI (Python)", "High-performance async API framework for REST endpoints"),
    ("Database", "MySQL + SQLAlchemy", "Relational database with ORM for structured data storage"),
    ("AI Engine", "OpenRouter API (Llama 3.1 8B)", "Large language model for resume analysis, career predictions, and chat"),
    ("Auth", "JWT (JSON Web Tokens)", "Stateless authentication for secure API access"),
    ("Resume Parsing", "NLP Libraries", "PDF/DOCX text extraction and skill entity recognition"),
    ("Charts", "Recharts", "Radial and bar charts for ATS score visualization"),
    ("Icons", "Lucide React", "Consistent, lightweight icon system"),
    ("HTTP Client", "Axios", "Promise-based API communication with request/response interceptors"),
    ("Version Control", "Git & GitHub", "Source code management and collaboration"),
]

table = doc.add_table(rows=1 + len(tech_data), cols=3)
table.style = "Light Grid Accent 1"

# Header row
hdr = table.rows[0].cells
for i, name in enumerate(["Layer", "Technology", "Purpose"]):
    hdr[i].text = name
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(10)

# Data rows
for r_idx, (layer, tech, purpose) in enumerate(tech_data, start=1):
    cells = table.rows[r_idx].cells
    cells[0].text = layer
    cells[1].text = tech
    cells[2].text = purpose
    for c_idx in range(3):
        for p in cells[c_idx].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

# ══════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. System Architecture", level=1)

add_body(
    "The application follows a clean three-tier architecture: a React-based "
    "single-page application (SPA) frontend communicates with a FastAPI backend "
    "over RESTful JSON APIs, which in turn interacts with a MySQL database through "
    "SQLAlchemy ORM. The AI capabilities are provided by an external large language "
    "model accessed via the OpenRouter API."
)

p = doc.add_paragraph()
run = p.add_run("Frontend Layer: ")
run.bold = True
p.add_run(
    "Built with React 19 and Vite, the SPA handles routing "
    "(react-router-dom), state management (React Context), and API communication "
    "(Axios with JWT interceptors). The UI uses a clean design system with CSS "
    "custom properties for consistent theming."
)

p = doc.add_paragraph()
run = p.add_run("Backend Layer: ")
run.bold = True
p.add_run(
    "FastAPI serves as the API gateway, handling authentication "
    "(JWT token generation and validation), resume file uploads with NLP-based "
    "parsing, and AI inference requests. The backend exposes endpoints for auth, "
    "users, resumes, jobs, courses, and AI operations."
)

p = doc.add_paragraph()
run = p.add_run("Data Layer: ")
run.bold = True
p.add_run(
    "MySQL stores user accounts, resume data, parsed analysis "
    "results, and cached recommendations. SQLAlchemy provides ORM abstraction "
    "with automatic table creation on startup."
)

# ══════════════════════════════════════════════════════════════
# 5. SCREENSHOTS AND OUTPUT
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. Screenshots and Output", level=1)
add_body("Below are screenshots demonstrating the key features of the CareerAI platform.")

add_image_with_caption(
    "localhost_5173_career (1).png",
    "Fig 1: Dashboard — Resume upload, ATS score overview, quick navigation, and resume history"
)

add_image_with_caption(
    "localhost_5173_career.png",
    "Fig 2: Career Intelligence — ATS score breakdown, detected skills, skill gap analysis, and career paths"
)

add_image_with_caption(
    "localhost_5173_career (4).png",
    "Fig 3: Job Recommendations — AI-curated job listings with match scores and skill tags"
)

add_image_with_caption(
    "localhost_5173_career (3).png",
    "Fig 4: Course Recommendations — Skill-building courses matched to identified gaps"
)

add_image_with_caption(
    "localhost_5173_career (2).png",
    "Fig 5: AI Career Coach — Conversational interface with resume context and quick prompts"
)

# ══════════════════════════════════════════════════════════════
# 6. KEY FEATURES SUMMARY
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. Key Features Summary", level=1)

features = [
    ("Resume Upload & Parsing", "Supports PDF, DOCX, and TXT formats. Uses NLP to extract skills, roles, experience, and education from unstructured resume text."),
    ("ATS Score Analysis", "Generates a compatibility score (0–100) against applicant tracking system benchmarks, displayed with a radial chart and progress bar."),
    ("Skill Gap Detection", "Compares detected skills against target role requirements. Each gap includes a priority level (high/medium/low) and an explanation."),
    ("Career Path Prediction", "Recommends career trajectories with match percentages, required skills, salary ranges, and estimated timelines."),
    ("Job Recommendations", "AI-curated job listings with match scores, company info, location, salary, job type, and direct apply links."),
    ("Course Recommendations", "Personalized learning resources targeting skill gaps, with match scores, ratings, duration, and provider info."),
    ("AI Career Coach", "Chatbot interface with resume context awareness, quick prompt suggestions, and conversation history."),
    ("Responsive Design", "Clean, professional UI that works across desktop and mobile viewports with a consistent design system."),
]

for title, desc in features:
    doc.add_heading(title, level=2)
    add_body(desc)

# ══════════════════════════════════════════════════════════════
# 7. CONCLUSION
# ══════════════════════════════════════════════════════════════
doc.add_heading("7. Conclusion", level=1)

add_body(
    "The CareerAI platform successfully delivers an integrated, AI-powered career "
    "intelligence system that addresses the key challenges job seekers face today. "
    "By combining resume parsing, ATS scoring, skill gap analysis, career path "
    "prediction, and an interactive AI coach into a single platform, it provides "
    "a comprehensive toolkit for career development."
)
add_body(
    "The project demonstrates the practical application of large language models "
    "in career services, showing how AI can transform unstructured resume data into "
    "structured, actionable insights. The clean architecture — separating concerns "
    "across React frontend, FastAPI backend, and MySQL database — ensures the "
    "system is maintainable and scalable."
)
add_body(
    "Future enhancements could include integration with live job boards via APIs, "
    "LinkedIn profile analysis, multi-language resume support, interview preparation "
    "modules, and collaborative features for career mentors and coaches."
)

# ══════════════════════════════════════════════════════════════
# 8. ACKNOWLEDGMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading("8. Acknowledgment", level=1)

add_body(
    "I would like to express my sincere gratitude to the Infosys Springboard "
    "team and my mentor for their continuous guidance and support throughout the "
    "development of this project. This virtual internship has been a valuable "
    "experience in learning full-stack development, AI integration, and building "
    "real-world applications that solve meaningful problems."
)
add_body(
    "The mentorship and structured project framework provided by the Infosys "
    "Springboard program were instrumental in shaping the direction and quality "
    "of this project. I am grateful for the opportunity to work on a project "
    "that combines artificial intelligence with practical career development tools."
)

# ─── Save ─────────────────────────────────────────────────────
doc.save(OUTPUT_DOCX)
print(f"DOCX generated: {OUTPUT_DOCX}")
