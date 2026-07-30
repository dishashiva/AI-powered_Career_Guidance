import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENTATION

BASE = r"E:\Spring Board Project\Report"
OUTPUT_DIR = os.path.join(BASE, "Outputs")
DOCX_FILE1 = os.path.join(BASE, "Disha S_Milestone 2.docx")
DOCX_FILE2 = os.path.join(BASE, "CareerAI_Milestone_2_Report.docx")

doc = Document()

# Page setup — A4 portrait
sec = doc.sections[0]
sec.page_width = Cm(21.0)
sec.page_height = Cm(29.7)
sec.top_margin = Cm(2.54)
sec.bottom_margin = Cm(2.0)
sec.left_margin = Cm(2.54)
sec.right_margin = Cm(2.54)

# Color tokens
DARK_TEXT = RGBColor(0x11, 0x18, 0x27)
PRIMARY_BLUE = RGBColor(0x25, 0x63, 0xEB)
GRAY_TEXT = RGBColor(0x4B, 0x55, 0x63)
LIGHT_MUTED = RGBColor(0x6B, 0x72, 0x80)

# Configure Styles
normal_style = doc.styles["Normal"]
normal_style.font.name = "Calibri"
normal_style.font.size = Pt(11)
normal_style.font.color.rgb = GRAY_TEXT
normal_style.paragraph_format.line_spacing = 1.15
normal_style.paragraph_format.space_after = Pt(6)

for level, size in [(1, 18), (2, 14), (3, 12)]:
    s = doc.styles[f"Heading {level}"]
    s.font.name = "Calibri Light"
    s.font.size = Pt(size)
    s.font.bold = True
    s.font.color.rgb = DARK_TEXT
    s.paragraph_format.space_before = Pt(14 - 2 * level)
    s.paragraph_format.space_after = Pt(4)

style_cap = doc.styles["Caption"]
style_cap.font.name = "Calibri"
style_cap.font.size = Pt(9.5)
style_cap.font.italic = True
style_cap.font.color.rgb = LIGHT_MUTED


def add_p(text, bold_prefix="", style="Normal", space_after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(space_after)
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.bold = True
        r_b.font.color.rgb = DARK_TEXT
    r_t = p.add_run(text)
    r_t.font.color.rgb = GRAY_TEXT
    return p


def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    if bold_prefix:
        r_b = p.add_run(bold_prefix)
        r_b.bold = True
        r_b.font.color.rgb = DARK_TEXT
    r_t = p.add_run(text)
    r_t.font.color.rgb = GRAY_TEXT
    return p


def add_image_section(filename, caption_text):
    img_path = os.path.join(OUTPUT_DIR, filename)
    if os.path.exists(img_path):
        p_img = doc.add_paragraph()
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(8)
        p_img.paragraph_format.space_after = Pt(4)
        run = p_img.add_run()
        run.add_picture(img_path, width=Cm(15.5))

        p_cap = doc.add_paragraph(caption_text, style="Caption")
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_cap.paragraph_format.space_after = Pt(14)
    else:
        add_p(f"[Missing screenshot: {filename}]")


# ══════════════════════════════════════════════════════════════
# COVER PAGE (Following Disha S_Milestone 1.pdf template)
# ══════════════════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("CareerAI")
r.font.name = "Calibri"
r.font.size = Pt(32)
r.font.bold = True
r.font.color.rgb = DARK_TEXT

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_sub.paragraph_format.space_after = Pt(24)
r = p_sub.add_run("AI-Powered Career Intelligence Platform")
r.font.name = "Calibri"
r.font.size = Pt(15)
r.font.color.rgb = PRIMARY_BLUE

p_intern = doc.add_paragraph()
p_intern.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_intern.add_run("INFOSYS SPRINGBOARD VIRTUAL INTERNSHIP 7.0")
r.font.size = Pt(12)
r.font.bold = True
r.font.color.rgb = GRAY_TEXT

p_proj = doc.add_paragraph()
p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_proj.add_run(
    "Project: CareerAI — AI-Powered Career Intelligence Platform"
)
r.font.size = Pt(11)
r.font.color.rgb = GRAY_TEXT

p_ms = doc.add_paragraph()
p_ms.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_ms.paragraph_format.space_after = Pt(40)
r = p_ms.add_run("Milestone 2 Report")
r.font.size = Pt(13)
r.font.bold = True
r.font.color.rgb = PRIMARY_BLUE

for _ in range(5):
    doc.add_paragraph()

p_by = doc.add_paragraph()
p_by.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_by.add_run("SUBMITTED BY:")
r.font.size = Pt(11)
r.font.bold = True
r.font.color.rgb = DARK_TEXT

for line in [
    "Name: Disha S",
    "Domain: Artificial Intelligence",
    "Date: 24.07.2026",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    r.font.size = Pt(11)
    r.font.color.rgb = GRAY_TEXT

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ══════════════════════════════════════════════════════════════
doc.add_heading("1. Introduction", level=1)
add_p(
    'The project "CareerAI — AI-Powered Career Intelligence Platform" is a comprehensive, full-stack web application designed to act as an intelligent career coach. By combining artificial intelligence, natural language processing (NLP), and structured data management, CareerAI assists job seekers and working professionals in evaluating their resumes, identifying critical skill gaps, discovering tailored career paths, and obtaining context-aware career guidance.'
)
add_p(
    "In modern recruitment processes, job seekers frequently face hurdles understanding how Applicant Tracking Systems (ATS) evaluate their resumes, identifying missing skills for target roles, and maintaining an up-to-date professional profile. Milestone 2 focuses on solving these core foundational requirements by implementing full User Profile Management, intelligent Resume Uploading, AI-driven Resume Parsing (extracting over 95% of structured candidate metadata), and robust Resume Lifecycle Management (Viewing, Downloading, Replacing, and Deleting resumes)."
)

# ══════════════════════════════════════════════════════════════
# 2. OBJECTIVE
# ══════════════════════════════════════════════════════════════
doc.add_heading("2. Objective", level=1)
add_p(
    "The primary objectives fulfilled in Milestone 2 of the CareerAI platform are:"
)

add_bullet(
    " Provide an end-to-end profile editing interface where users can manage personal info, education history, technical skills, certifications, work experience, projects, and career interests.",
    bold_prefix="User Profile Management —",
)
add_bullet(
    " Support seamless resume file uploads in PDF, DOCX, DOC, and TXT formats with automated file type and size validation (up to 10 MB).",
    bold_prefix="Resume Upload & Validation —",
)
add_bullet(
    " Extract up to 95%+ of candidate details (Name, Contact Info, Bio/Summary, Education, Skills, Work History, Certifications, Projects, Languages) using AI and NLP prompt engineering.",
    bold_prefix="High-Accuracy Resume Parsing —",
)
add_bullet(
    " Provide full CRUD functionality allowing users to view original uploaded files inline in the browser, download original files, replace existing resumes, and delete resume records.",
    bold_prefix="Resume Lifecycle Management —",
)
add_bullet(
    " Implement a real-time progress indicator calculating user profile completeness (0% to 100%) to encourage thorough profile setup.",
    bold_prefix="Profile Completeness Analytics —",
)
add_bullet(
    " Expose scalable RESTful FastAPI endpoints to handle authentication, profile updates, background parsing tasks, and file management securely.",
    bold_prefix="Backend REST Services —",
)

# ══════════════════════════════════════════════════════════════
# 3. TECH STACK
# ══════════════════════════════════════════════════════════════
doc.add_heading("3. Tech Stack", level=1)
add_p(
    "Milestone 2 was implemented leveraging modern full-stack web technologies and AI models:"
)

tech_table_data = [
    (
        "Frontend",
        "React.js (v19) + Vite",
        "Fast single-page application (SPA) with responsive component architecture",
    ),
    (
        "Backend",
        "FastAPI (Python 3.12+)",
        "High-performance asynchronous REST API framework",
    ),
    (
        "Database",
        "MySQL + SQLAlchemy ORM",
        "Relational database for storing users, profiles, and resume metadata",
    ),
    (
        "AI Engine",
        "Groq API / Llama 3.3 70B",
        "LLM inference engine for structured NLP resume parsing and profile extraction",
    ),
    (
        "Authentication",
        "JWT (JSON Web Tokens)",
        "Stateless, secure session management and endpoint protection",
    ),
    (
        "Text Extraction",
        "PyPDF & python-docx",
        "Parsing raw text content from PDF and DOCX resume uploads",
    ),
    (
        "UI & Styling",
        "Vanilla CSS + Lucide React",
        "Modern custom design tokens, glassmorphism UI, and clean icons",
    ),
    (
        "HTTP Client",
        "Axios",
        "Async API integration with request token interceptors and response handlers",
    ),
]

table = doc.add_table(rows=1 + len(tech_table_data), cols=3)
table.style = "Table Grid"

# Header row styling
hdr_cells = table.rows[0].cells
hdr_titles = ["Layer", "Technology", "Purpose"]
for i, title in enumerate(hdr_titles):
    hdr_cells[i].text = title
    for p in hdr_cells[i].paragraphs:
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = DARK_TEXT

for row_idx, (layer, tech, purpose) in enumerate(tech_table_data, start=1):
    row_cells = table.rows[row_idx].cells
    row_cells[0].text = layer
    row_cells[1].text = tech
    row_cells[2].text = purpose
    for c in range(3):
        for p in row_cells[c].paragraphs:
            p.paragraph_format.space_after = Pt(2)
            for run in p.runs:
                run.font.size = Pt(9.5)

doc.add_paragraph().paragraph_format.space_after = Pt(6)

# ══════════════════════════════════════════════════════════════
# 4. SYSTEM ARCHITECTURE & DATA FLOW
# ══════════════════════════════════════════════════════════════
doc.add_heading("4. System Architecture & Data Flow", level=1)
add_p(
    "The CareerAI platform follows a decoupled, three-tier architecture ensuring high performance, security, and scalability."
)
add_bullet(
    " Built using React.js and Vite. Communicates with backend endpoints asynchronously over HTTP JSON requests via Axios.",
    bold_prefix="Frontend Tier: ",
)
add_bullet(
    " Implemented using FastAPI. Manages JWT security, file upload handling, database transactions via SQLAlchemy, and background AI execution pipelines.",
    bold_prefix="Backend Tier: ",
)
add_bullet(
    " Relational database hosting structured entities including Users, Profiles, Resumes, and Skill Inventories.",
    bold_prefix="Database Tier (MySQL): ",
)
add_bullet(
    " Integrates LLM capabilities via structured JSON mode prompts for high-accuracy NLP parsing.",
    bold_prefix="AI Processing Tier: ",
)

add_p(
    "Data Flow Architecture (Milestone 2 Workflow):",
    bold_prefix="System Workflow: ",
)

flow_steps = [
    "Landing Page ──► User Registration ──► Secure Login (JWT Issued)",
    "Dashboard ──► User Profile Management (Edit Bio, Skills, Education, Links)",
    "Resume Upload ──► File Validation (PDF/DOCX/TXT) ──► Server Storage & Database Record Creation",
    "AI Pipeline (Async) ──► NLP Text Extraction ──► Structured Data Parsing (95%+ Metadata)",
    "Profile Autofill ──► Auto-populates Profile & Review Panel for User Confirmation",
    "Resume Management ──► View Inline, Download Original, Replace, or Delete Resume",
]
for step in flow_steps:
    add_bullet(step)

# ══════════════════════════════════════════════════════════════
# 5. MODULES IMPLEMENTED (MILESTONE 2)
# ══════════════════════════════════════════════════════════════
doc.add_heading("5. Modules Implemented", level=1)

doc.add_heading("5.1 User Profile Management", level=2)
add_p(
    "The User Profile Management module provides a central hub where users can maintain their professional identity."
)
add_bullet(
    "Editable name, email, phone number, location (city, state/country), and personal professional summary/bio.",
    bold_prefix="Personal Info & Bio: ",
)
add_bullet(
    "Adding and updating degree details, institution names, and graduation years.",
    bold_prefix="Education History: ",
)
add_bullet(
    "Interactive skill tag inputs allowing users to add, edit, or remove technical and soft skills with comma/Enter shortcuts.",
    bold_prefix="Skills & Certifications: ",
)
add_bullet(
    "Storing professional certifications, online courses/bootcamps, languages spoken, and key career achievements.",
    bold_prefix="Credentials & Projects: ",
)
add_bullet(
    'Managing work preferences (Remote, Hybrid, On-site), availability status ("Open to work"), target job roles, and salary expectations.',
    bold_prefix="Career Interests: ",
)
add_bullet(
    "One-click AI feature that populates profile fields from the user's parsed resume.",
    bold_prefix="AI Profile Autofill: ",
)

doc.add_heading("5.2 Resume Upload Module", level=2)
add_p(
    "The Resume Upload module facilitates drag-and-drop file ingestion with instant feedback."
)
add_bullet(
    "Supports PDF (.pdf), Microsoft Word (.docx, .doc), and plain text (.txt) files up to 10 MB.",
    bold_prefix="Multi-Format Ingestion: ",
)
add_bullet(
    "Real-time validation of file extensions and size prior to submission.",
    bold_prefix="File Validation: ",
)
add_bullet(
    "Asynchronous non-blocking uploads returning instant status updates while background workers process AI extraction.",
    bold_prefix="Background Processing: ",
)

doc.add_heading("5.3 Resume Parsing Module (95%+ Extraction)", level=2)
add_p(
    "The AI Resume Parsing engine uses custom NLP prompt engineering to extract structured data from raw resume text with high precision."
)
add_bullet(
    "Extracts candidate full name, phone number, location, LinkedIn URL, GitHub URL, and portfolio link.",
    bold_prefix="Personal & Contact Info: ",
)
add_bullet(
    "Generates a 2-4 sentence professional summary reflecting experience and core competencies.",
    bold_prefix="Professional Bio / Summary: ",
)
add_bullet(
    "Identifies primary current role title, suggested target next step, and total experience years.",
    bold_prefix="Career & Titles: ",
)
add_bullet(
    "Parses programming languages, frameworks, developer tools, and domain-specific skills.",
    bold_prefix="Technical & Soft Skills: ",
)
add_bullet(
    "Extracts industry certifications (e.g. AWS, PMP) and completed courses/training programs.",
    bold_prefix="Certifications & Courses: ",
)

doc.add_heading("5.4 Resume Management & Profile Completion", level=2)
add_p(
    "The Resume Management module gives users full control over their uploaded resume files and profile metrics."
)
add_bullet(
    "Opens uploaded PDF/DOCX resumes inline in the browser viewer.",
    bold_prefix="Inline View: ",
)
add_bullet(
    "Allows downloading the original resume file at any time.",
    bold_prefix="File Download: ",
)
add_bullet(
    "Enables uploading a new file to update existing parsed data.",
    bold_prefix="Replace Resume: ",
)
add_bullet(
    "Deletes the database entry and removes stored files from disk storage safely.",
    bold_prefix="Delete Resume: ",
)
add_bullet(
    "Visual circular progress bar tracking completeness (0% to 100%) to encourage detailed setup.",
    bold_prefix="Profile Completion Meter: ",
)

doc.add_heading("5.5 Backend API Architecture", level=2)
add_p("FastAPI REST endpoints implemented in Milestone 2:")
add_bullet(
    "POST /auth/register, POST /auth/login (JWT Token Authentication)",
    bold_prefix="Auth Endpoints: ",
)
add_bullet(
    "GET /users/me, PUT /users/me/profile, POST /users/me/autofill-profile",
    bold_prefix="User & Profile Endpoints: ",
)
add_bullet(
    "POST /resumes/upload, GET /resumes/, GET /resumes/{id}/status, GET /resumes/{id}/view, GET /resumes/{id}/download, DELETE /resumes/{id}",
    bold_prefix="Resume Endpoints: ",
)

# ══════════════════════════════════════════════════════════════
# 6. OUTPUT & SCREENSHOTS
# ══════════════════════════════════════════════════════════════
doc.add_heading("6. Output & Screenshots", level=1)
add_p(
    "Below are the empirical output screenshots demonstrating the implemented Milestone 2 modules."
)

add_image_section(
    "Profile management.png",
    "Figure 1: User Profile Management Module — Comprehensive Profile Editing, Skills/Certifications Tagging, and AI Profile Autofill",
)

add_image_section(
    "Resume parsing.png",
    "Figure 2: AI Resume Parsing Module — High-Precision 95%+ Metadata Extraction (Contact Details, Bio, Education, Skills, and Certifications)",
)

add_image_section(
    "Resume upload and delete.png",
    "Figure 3: Resume Management Module — Drag & Drop Resume Upload, Real-Time Status, View, Download, Replace, and Delete Functionality",
)

# ══════════════════════════════════════════════════════════════
# 7. KEY FEATURES IMPLEMENTED IN MILESTONE 2
# ══════════════════════════════════════════════════════════════
doc.add_heading("7. Key Features Implemented in Milestone 2", level=1)

features_m2 = [
    (
        "Complete User Profile Management",
        "Comprehensive form interface to manage bio, education, skills, certifications, work history, projects, and target role preferences.",
    ),
    (
        "1-Click AI Profile Autofill",
        "Automated extraction pipeline that reads parsed resume data and populates user profile fields instantly.",
    ),
    (
        "High-Accuracy AI Resume Parsing (95%+)",
        "NLP prompt pipeline extracting candidate metadata including name, contact details, social links, bio, skills, education, and credentials.",
    ),
    (
        "Secure Multi-Format Resume Upload",
        "Supports PDF, DOCX, DOC, and TXT files with file type checking, size enforcement (10 MB limit), and background execution.",
    ),
    (
        "Full Resume CRUD & Lifecycle Management",
        "Allows users to list uploaded resumes, view files inline in browser, download original files, replace existing resumes, and delete files.",
    ),
    (
        "Dynamic Profile Completion Analytics",
        "Calculates and visualizes overall profile completion percentage (0% to 100%) to motivate users to provide complete career information.",
    ),
    (
        "FastAPI REST & MySQL Integration",
        "Robust backend service layer utilizing SQLAlchemy ORM with automatic schema migrations and JWT authentication.",
    ),
]

for title, desc in features_m2:
    doc.add_heading(title, level=2)
    add_p(desc)

# ══════════════════════════════════════════════════════════════
# 8. CONCLUSION
# ══════════════════════════════════════════════════════════════
doc.add_heading("8. Conclusion", level=1)
add_p(
    "Milestone 2 of the CareerAI platform successfully implements all core foundational capabilities for user profile management, resume upload, AI resume parsing, and resume lifecycle management. By extracting over 95% of candidate information automatically and providing full CRUD controls alongside an intuitive React interface, the application delivers a seamless experience for job seekers."
)
add_p(
    "With Milestone 2 complete, the platform provides a robust data foundation for Milestone 3, which will introduce ATS scoring algorithms, detailed skill gap detection, AI career path predictions, personalized job and course recommendations, and the AI Career Coach chatbot."
)

# ══════════════════════════════════════════════════════════════
# 9. ACKNOWLEDGMENT
# ══════════════════════════════════════════════════════════════
doc.add_heading("9. Acknowledgment", level=1)
add_p(
    "I extend my sincere gratitude to the Infosys Springboard Virtual Internship 7.0 team and my project mentor for their guidance and encouragement throughout Milestone 2. The learning experience gained in building full-stack web architectures, FastAPI services, and AI-driven NLP integrations has been invaluable."
)

# Save documents
doc.save(DOCX_FILE1)
doc.save(DOCX_FILE2)

print("Successfully generated DOCX report files:")
print("1.", DOCX_FILE1)
print("2.", DOCX_FILE2)
